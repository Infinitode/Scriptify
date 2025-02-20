import random
import string
import webview
import os
import time
import sys
import threading
import whisper  # Whisper for speech-to-text processing
import queue
import json  # Import json for safe string handling
import sounddevice as sd  # For audio recording
import numpy as np
from scipy.io.wavfile import write  # To handle audio format conversion
from docx import Document
from docx.shared import Pt
from fpdf import FPDF
from datetime import datetime
from bs4 import BeautifulSoup
from whisper.utils import get_writer

# Transcription queue for thread communication
transcription_queue = queue.Queue()
is_transcribing = False
transcribed_text = ""  # Global variable to hold the transcribed text

model = None

# Adjust base path for bundled applications
if getattr(sys, 'frozen', False):
    base_path = sys._MEIPASS
else:
    base_path = os.path.abspath(".")

# Set whisper download folder
os.environ["WHISPER_DOWNLOAD_DIR"] = os.path.join(base_path, "models")
os.environ["WHISPER_CACHE_DIR"] = os.path.join(base_path, "models")

# Create if it doesn't exist
if not os.path.exists(os.environ["WHISPER_DOWNLOAD_DIR"]):
    os.makedirs(os.environ["WHISPER_DOWNLOAD_DIR"])

# Set the path for the WAV file
wav_file_path = os.path.join(base_path, "temp_audio.wav")

def generate_random_string(length=6):
    characters = string.ascii_letters + string.digits
    return ''.join(random.choices(characters, k=length))


def randomName(type, amount, extension):
    random_str = generate_random_string()
    return f"{type}_{amount}_samples_{random_str}{extension}"


def escape_java_script_string(text):
    # Use json.dumps to escape special characters for JavaScript
    return json.dumps(text)  # This handles quotes, newlines, etc.


class Api:
    def __init__(self):
        self.is_fullscreen = False
        self._window = None
        self.transcribing = False
        self.timer_running = False  # Add this flag to control the timer

    def set_window(self, window):
        self._window = window

    def closeWindow(self):
        self._window.destroy()

    def minimizeWindow(self):
        self._window.minimize()

    def toggleFullscreen(self):
        self.is_fullscreen = not self.is_fullscreen
        self._window.toggle_fullscreen()

    def download_model(self, model_name):
        global model
        try:
            model = whisper.load_model(model_name, download_root=os.environ["WHISPER_DOWNLOAD_DIR"])
            launch_application()
        except Exception as e:
            print(f"Error loading model: {e}")

    def display_text(self, text):
        self._window.evaluate_js(f"displayText({escape_java_script_string(text)})")

    def update_progress_bar(self, progress):
        """
        Updates the progress bar in the webview.
        """
        # Update the progress bar element in the webview using JavaScript
        self._window.evaluate_js(f"updateProgressBar({progress})")

    # Start transcription process
    def start_transcription(self):
        global is_transcribing
        if not is_transcribing:
            start_continuous_transcription()
            return "Transcription started"
        return "Transcription is already running."

    # Stop transcription process
    def stop_transcription(self):
        global is_transcribing
        if is_transcribing:
            stop_continuous_transcription()
            api.stop_timer()
            return "Transcription stopped."
        return "Transcription is not running."

    def update_transcribed_text(self, text):
        global transcribed_text
        transcribed_text = text  # Update the global variable

        # Escape the text before sending it to JavaScript
        escaped_text = escape_java_script_string(text)
        self._window.evaluate_js(f"updateTranscribedText({escaped_text})")

    def export_as(self):
        # Show save dialog and get the file path, including the selected format
        file_path = self._window.create_file_dialog(webview.SAVE_DIALOG, file_types=[
            ("JSON File (*.json)"),
            ("Markdown File (*.md)"),
            ("Plain Text File (*.txt)"),
            ("HTML File (*.html)"),
            ('Word Document (*.docx)'),
            ('PDF (*.pdf)')
        ])

        if file_path:
            try:
                # Get the content of <h1> and <p> elements from the webview
                title_html = self._window.evaluate_js(
                    'document.querySelector(".text h1").innerHTML')
                content_html = self._window.evaluate_js(
                    'document.querySelector(".text p").innerHTML')

                # Use BeautifulSoup to handle the HTML content and preserve formatting
                title = self._parse_html_to_text(title_html)
                content = self._parse_html_to_text(content_html)

                # Determine the file format based on the file extension
                file_extension = os.path.splitext(file_path)[1].lower()

                # Prepare and save the data based on the chosen format
                if file_extension == ".json":
                    data = {
                        "title": title,
                        "content": content
                    }
                    with open(file_path, 'w', encoding='utf-8') as file:
                        json.dump(data, file, indent=1)

                elif file_extension == ".md":
                    data = f"# {title}\n\n{content}"
                    with open(file_path, 'w', encoding='utf-8') as file:
                        file.write(data)

                elif file_extension == ".txt":
                    data = f"{title}\n\n{content}"
                    with open(file_path, 'w', encoding='utf-8') as file:
                        file.write(data)

                elif file_extension == ".html":
                    data = f"<h1>{title}</h1>\n<p>{content}</p>"
                    with open(file_path, 'w', encoding='utf-8') as file:
                        file.write(data)

                elif file_extension == ".docx":
                    self._export_to_word(file_path, title, content)

                elif file_extension == ".pdf":
                    self._export_to_pdf(file_path, title, content)

                return "Content saved successfully."
            except Exception as e:
                return f"Error saving content: {e}"

    # Method to save content to a .scriptify file
    def save_content(self):
        # Use pywebview to open the file save dialog
        file_path = self._window.create_file_dialog(webview.SAVE_DIALOG, file_types=[
                                                    'Scriptify File (*.scriptify)'])

        if file_path:
            try:
                # Get the content of <h1> and <p> elements from the webview
                title = self._window.evaluate_js(
                    'document.querySelector(".text h1").innerHTML')
                content = self._window.evaluate_js(
                    'document.querySelector(".text p").innerHTML')

                # Create a dictionary with the title and content
                data = {
                    "title": title,
                    "content": content
                }

                # Save the data as JSON to the selected file
                with open(file_path, 'w', encoding='utf-8') as file:
                    json.dump(data, file, indent=4)

                return "Content saved successfully."
            except Exception as e:
                return f"Error saving content: {e}"

    def auto_backup(self):
        try:
            # Get the content of <h1> and <p> elements from the webview
            title_html = self._window.evaluate_js(
                'document.querySelector(".text h1").innerHTML')
            content_html = self._window.evaluate_js(
                'document.querySelector(".text p").innerHTML')
            # Use BeautifulSoup to handle the HTML content and preserve formatting
            title = self._parse_html_to_text(title_html)
            content = self._parse_html_to_text(content_html)

            # Sanitize the title for the backup filename
            sanitized_title = title.strip().replace(" ", "_")  # Replace spaces with underscores

            # Create the backup filename based on the title and timestamp
            # Get the current date and time
            now = datetime.now()

            # Format the date and time to fit in a filename
            filename_time = now.strftime("%Y-%m-%d_%H-%M-%S")
            backup_filename = f"{sanitized_title}_{filename_time}_backup.scriptify"
            backup_path = os.path.join(base_path, backup_filename)

            # Prepare the data to be saved
            data = {
                "title": title,
                "content": content
            }

            # Save the data to the file as JSON
            with open(backup_path, 'w', encoding='utf-8') as file:
                json.dump(data, file, indent=4)

            print("Saved to:", backup_path)

            return f"Auto backup saved to {backup_filename}."

        except Exception as e:
            return f"Error during auto backup: {e}"

    # Method to load content from a .scriptify file
    def load_content(self):
        # Use pywebview to open the file open dialog
        file_paths = self._window.create_file_dialog(webview.OPEN_DIALOG, file_types=[
                                                     'Scriptify File (*.scriptify)'], allow_multiple=False)

        if file_paths and len(file_paths) > 0:
            file_path = file_paths[0]
            # Debugging print to check file selection
            print(f"Selected file: {file_path}")

            try:
                # Open and load the JSON file
                with open(file_path, 'r', encoding='utf-8') as file:
                    data = json.load(file)

                # Extract title and content
                title = data.get("title", "")
                content = data.get("content", "")

                # Use JSON.stringify to safely inject text content in JavaScript
                self._window.evaluate_js(f'''
                    document.querySelector(".text h1").innerHTML = JSON.stringify({json.dumps(title)}).replace(/^"|"$/g, '');
                    document.querySelector("#transcribedText").innerHTML = JSON.stringify({json.dumps(content)}).replace(/^"|"$/g, '');
                ''')

                return "Content loaded successfully."
            except Exception as e:
                return f"Error loading content: {e}"

    def set_icon(self, icon):
        stringy = ""
        if icon == 1:
            stringy = "document.getElementById('transcribe').innerHTML = '<i class=\"bi bi-mic-fill\"></i>'"
            stringy2 = "document.getElementById('timer').innerHTML = 'Processing speech, please wait...'"
            self._window.evaluate_js(stringy)
            self._window.evaluate_js(stringy2)
        elif icon == 2:
            stringy = "document.getElementById('transcribe').innerHTML = '<i class=\"bi bi-mic-mute-fill color\"></i>'"
            self._window.evaluate_js(stringy)

    def update_timer(self, remaining_time):
        self._window.evaluate_js(f"updateTimer({remaining_time})")

    # Timer function
    def timer(self, duration):
        self.timer_running = True  # Start the timer
        for remaining_time in range(duration, 0, -1):
            if not self.timer_running:
                break  # Stop the timer immediately if the flag is set to False
            self.update_timer(remaining_time)
            time.sleep(1)
        self.update_timer(0)  # Reset the timer when done or stopped

    # Stop the timer by setting the flag to False
    def stop_timer(self):
        self.timer_running = False  # Ensure this stops the timer
        self.update_timer(0)  # Reset the UI to show the timer has stopped

    def _parse_html_to_text(self, html):
        """
        Parse HTML and convert it to plain text while preserving whitespace, line breaks, etc.
        """
        soup = BeautifulSoup(html, 'html.parser')

        # Replace <br> and <p> tags with line breaks to preserve structure
        for br in soup.find_all("br"):
            br.replace_with("\n")
        for p in soup.find_all("p"):
            p.insert_before("\n")
            p.append("\n")
        for h in soup.find_all("h2"):
            h.insert_before("\n")
            h.append("\n")

        # Return text with preserved whitespace and structure
        return soup.get_text()

    def _export_to_word(self, file_path, title, content):
        # Create a Word document
        doc = Document()

        # Add title with sans-serif font (General Sans)
        title_paragraph = doc.add_paragraph(title)
        title_run = title_paragraph.runs[0]
        title_run.font.name = 'Arial'  # Set the font for the title
        title_run.font.size = Pt(24)

        # Add content with serif font (Erode)
        content_paragraph = doc.add_paragraph(content)
        content_run = content_paragraph.runs[0]
        content_run.font.name = 'Times'  # Set the font for the content
        content_run.font.size = Pt(12)

        # Save the document
        doc.save(file_path)

    def _export_to_pdf(self, file_path, title, content):
        # Create a PDF document
        pdf = FPDF()
        pdf.add_page()

        # Add title with larger font size and sans-serif (PDF only supports built-in fonts)
        pdf.set_font("Arial", "B", 16)  # Simulating sans-serif using Arial for title
        pdf.cell(200, 10, txt=title, ln=True, align="C")

        # Add content with serif font (PDF only supports built-in fonts)
        pdf.set_font("Times", size=12)  # Simulating serif using Times for content
        pdf.multi_cell(0, 10, content)

        # Save the PDF
        pdf.output(file_path)

    def notify(self, content):
        self._window.evaluate_js(f"displayMessage({content})")

def launch_application():
    window = webview.create_window('Scriptify', 'index.html', js_api=api, width=1280, height=800,
                                   resizable=True, min_size=(1000, 600), background_color='#ffffff', frameless=True, easy_drag=False)
    api.set_window(window)
    webview.windows[0].destroy()

def record_audio_continuously(queue, sample_rate=16000, channels=1):
    """
    Continuously records audio in chunks and adds them to the queue for transcription.
    """
    chunk_duration = 10  # Duration of each audio chunk in seconds
    while is_transcribing:
        print("Recording audio chunk...")
        audio_data = sd.rec(int(chunk_duration * sample_rate), samplerate=sample_rate, channels=channels, dtype='float32')
        sd.wait()  # Wait for the recording to finish

        # Convert to int16 and add to queue
        audio_chunk = (audio_data * 32767).astype(np.int16)
        queue.put(audio_chunk)

        print("Audio chunk added to queue for transcription.")


def transcription_worker():
    """
    Worker that transcribes audio chunks from the queue using the Whisper model.
    """
    temp_file = os.path.join(base_path, "temp_audio_chunk.wav")  # Temporary file path

    while True:
        audio_chunk = transcription_queue.get()  # Wait for the next audio chunk
        if audio_chunk is None:  # Exit signal
            break

        print("Processing audio chunk for transcription...")
        write(temp_file, 16000, audio_chunk)  # Save chunk to WAV file

        # Transcribe audio chunk
        try:
            result = model.transcribe(temp_file, fp16=False)  # Set fp16=False for CPU
            chunk_text = result.get("text", "")
            print(f"Transcribed Chunk: {chunk_text}")

            # Update the global transcribed text and UI
            global transcribed_text
            transcribed_text = chunk_text + "\n"
            api.update_transcribed_text(transcribed_text)
        except Exception as e:
            print(f"Error in transcription: {e}")


def start_continuous_transcription():
    """
    Starts the continuous audio recording and transcription process.
    """
    global is_transcribing
    if is_transcribing:
        print("Transcription is already running.")
        return

    is_transcribing = True
    print("Starting continuous transcription...")

    # Start recording thread
    recording_thread = threading.Thread(
        target=record_audio_continuously, args=(transcription_queue,), daemon=True)
    recording_thread.start()

    # Start transcription thread
    transcription_thread = threading.Thread(target=transcription_worker, daemon=True)
    transcription_thread.start()


def stop_continuous_transcription():
    """
    Stops the continuous transcription process.
    """
    global is_transcribing
    if not is_transcribing:
        print("Transcription is not running.")
        return

    is_transcribing = False
    transcription_queue.put(None)  # Signal transcription worker to exit
    print("Continuous transcription stopped.")

# def transcription_worker():
#     model = whisper.load_model("base.en")  # Load the Whisper model once
#     # transcribe_file(model, 'test.wav', 'transcription.srt')
#     while True:
#         transcribing = transcription_queue.get()
#         if transcribing:
#             print("Transcription in progress...")
#             # Record and transcribe audio in real time
#             record_audio_and_transcribe(model)
#         else:
#             print("Transcription stopped.")
#             api.stop_timer()

def run_transcription_thread():
    thread = threading.Thread(target=transcription_worker, daemon=True)
    thread.start()


if __name__ == '__main__':
    model = None
    api = Api()
    window = webview.create_window('Scriptify Startup', 'launch.html', js_api=api, width=800, height=600,
                                   resizable=True, min_size=(800, 600), background_color='#ffffff', frameless=True, easy_drag=False)
    api.set_window(window)
    webview.start()